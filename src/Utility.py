from docx import Document
from docx.shared import RGBColor
from docx2pdf import convert
import fitz
import http.server
import os
import re
import socketserver
import threading
import tiktoken
import webbrowser
# from pdf2docx import Converter
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.chat_models import AzureChatOpenAI
from langchain.llms import AzureOpenAI
from langchain.embeddings import OpenAIEmbeddings
# 
import pysqlite3
import sys
sys.modules["sqlite3"] = sys.modules.pop("pysqlite3")
# import chromadb
from langchain.vectorstores import Chroma

footer = """Affine Inc. 2018 156th Avenue, N.E, Building F, Suite 333, Bellevue, Washington, 98007 Tel: +91 -80-6569 -0996 | Web: www.affine.ai | Mail: info@affine.ai Affine Confidential"""
special_characters = r"[]{}()^$.*+?|\\"
# Escape special characters by adding a backslash before them
escaped_string = re.sub(f"[{''.join(re.escape(char) for char in special_characters)}]", r"\\\g<0>", footer)
pattern = re.sub(r'\s+', r'\\s*', escaped_string)
replacement = " footer "


class TextProcessingUtility:
    def __init__(self):
        """
        Initializes an instance of the TextProcessingUtility class.

        This class provides a set of utility methods for processing and manipulating text data,
        including reading from PDF and DOCX files, highlighting text in PDFs, calculating OpenAI API costs,
        and various text-related operations.

        Note: Some methods in this class may rely on external classes or libraries for specific functionalities.

        Method Capabilities:
        - read_pdf(file_path): Reads the content of a PDF file and performs text processing.
        - read_docx_page_by_page(file_path): Reads the content of a DOCX file page by page and performs text processing.
        - split_text_newlines_regex(text): Splits text based on newline characters and performs text processing.
        - split_text_spliter(documents, user_input_chunk_size, user_input_chunk_overlap):
          Splits text into chunks using a specified text splitter and parameters.
        - model_load(deployment_name, model_type): Loads an OpenAI model based on deployment and model type.
        - highlight_text_in_pdf(pdf_path, final_path, text_to_highlight, single_change_highlights):
          Highlights specified text in a PDF file and saves the highlighted versions.
        - link_open(Handler, PORT): Starts a server for hosting files with specified handler and port.
        - replace_newlines(text): Replaces consecutive newlines with double newlines in the text.
        - download_docx(final_text, output_file_path): Downloads a DOCX file with specified content.
        - pdf_docx_converter(pdf_file_path, docx_file_dir): Converts a PDF file to DOCX format.
        - saved_updated_docx(updated_final_results, docx_file_path, new_docx_file_path):
          Replaces specified text in a DOCX file and saves the updated document.
        - updated_text_with_color(new_updated_file, updated_final_results):
          Highlights specified text in a DOCX file with a specified color and saves the updated document.
        - openai_api_calculate_cost(usage, model): Calculates the cost of OpenAI API usage based on usage data and model.
        - num_tokens_from_string(string, encoding_name): Returns the number of tokens in a text string.
        - docx_pdf(docx_file_path, pdf_file_name): Converts a DOCX file to a PDF file.

        Example usage:
        ```
        text_processor = TextProcessingUtility()
        ```
        """
        pass

    def read_pdf(self,file_path:str)-> str:
        """
        Reads text content from a PDF file and performs text processing.

        Parameters:
        - file_path (str): The path to the PDF file.

        Returns:
        - str: Processed text content from the PDF.
        """
        # Open the PDF file using PyMuPDF (fitz) library
        doc = fitz.open(file_path)
        # Print a message indicating that the PDF has been loaded
        print('PDF LOADED')
        # Initialize an empty string to store the combined text content from all pages
        documents = ""
        # Iterate through each page in the PDF document
        for page in doc:
            # Concatenate the text content of each page to the 'documents' string
            documents += ''.join(page.get_text())
        # Perform text processing using regular expressions (assuming 'pattern' and 'replacement' are defined)
        documents = re.sub(pattern, replacement, documents)
        # Return the processed text content
        return documents

    def read_docx_page_by_page(self,file_path:str)-> str:
        """
        Reads text content from a DOCX file page by page and performs text processing.

        Parameters:
        - file_path (str): The path to the DOCX file.

        Returns:
        - str: Processed text content from the DOCX file.
        """
        # Open the DOCX file using python-docx library
        doc = Document(file_path)
        # Initialize an empty string to store the combined text content from all paragraphs
        documents = ""
        # Print the total number of paragraphs in the document
        print("Total Paragraphs ::", len(doc.paragraphs))

        # Iterate through each paragraph in the DOCX document
        for n, paragraph in enumerate(doc.paragraphs):
            # Get the text content of each paragraph
            text = paragraph.text
            # Concatenate the text content of each paragraph to the 'documents' string
            documents += text + "\n"
        # Perform text processing using regular expressions (assuming 'pattern' and 'replacement' are defined)
        documents = re.sub(pattern, replacement, documents)
        # Return the processed text content
        return documents

    def split_text_newlines_regex(self,text:str) -> str:
        """
        Splits a text into a list of paragraphs using a regular expression based on consecutive newlines.

        Parameters:
        - text (str): The input text to be split.

        Returns:
        - list of str: List of paragraphs extracted from the input text.
        """
        # Use a regular expression to split the text based on consecutive newlines and optional spaces
        # The regular expression '(\n\s*){2,}' matches two or more consecutive newlines with optional spaces
        # The resulting list comprehension filters out empty strings and strings containing "footer"
        text = [text for text in re.split('(\n\s*){2,}', text) if text != "\n" and text != "footer  "]
        # Return the list of paragraphs
        return text

  
    def split_text_spliter(self,documents:str, user_input_chunk_size:int=2048, user_input_chunk_overlap:int=20) -> list[str]:
        """
        Splits a document into chunks using a RecursiveCharacterTextSplitter.

        Parameters:
        - documents (str): The document text to be split into chunks.
        - user_input_chunk_size (int): The desired size of each chunk (default is 2048 characters).
        - user_input_chunk_overlap (int): The desired overlap between consecutive chunks (default is 20 characters).

        Returns:
        - list of str: List of document chunks obtained by splitting the input text.
        """
        # Create an instance of the RecursiveCharacterTextSplitter class
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=user_input_chunk_size,
                                                       chunk_overlap=user_input_chunk_overlap,
                                                       length_function=len)
        # Use the text splitter to split the document into chunks
        document_chunks = text_splitter.split_text(documents)
        # Return the list of document chunks
        return document_chunks

    def model_load(self, deployment_name: (None, str) = None, model_type: str = "gpt-4"):
        """
        Load and initialize an OpenAI language model based on the specified model type.

        Parameters:
        - deployment_name (str): The deployment name for the language model (default is None).
        - model_type (str): The type of language model to load ("gpt-4" or any other type supported, default is "gpt-4").

        Returns:
        - AzureChatOpenAI or AzureOpenAI: An instance of the specified OpenAI language model.
        """
        if model_type == "gpt-4":
            print("GPT-4 Loaded")
            # Azure Open Ai GPT-4
            gpt4_llm = AzureChatOpenAI(temperature=0,
                                       top_p=0.1,
                                       deployment_name="gpt4-8k",
                                       model_kwargs={
                                           "api_key": os.environ["OPENAI_API_KEY"],
                                           "api_base": os.environ["OPENAI_API_BASE"],
                                           "api_type": os.environ["OPENAI_API_TYPE"],
                                           "api_version": os.environ["OPENAI_API_VERSION"],
                                       }
                                       )
            return gpt4_llm
        else:
            print("GPT-3.5 Loaded")
            # Azure Open Ai GPT-3.5
            gpt35_llm = AzureOpenAI(temperature=0,
                                     top_p=0.1,
                                     deployment_name="chatgpt",
                                     model_kwargs={
                                         "api_key": os.environ["OPENAI_API_KEY"],
                                         "api_base": os.environ["OPENAI_API_BASE"],
                                         "api_type": os.environ["OPENAI_API_TYPE"],
                                         "api_version": os.environ["OPENAI_API_VERSION"],
                                     }
                                     )
            return gpt35_llm

    def highlight_text_in_pdf(self,pdf_path:str, final_path:str, text_to_highlight:list[str], single_change_highlights:bool=True) -> list[str]:
        """
        Highlight specified text in a PDF document and save the result.

        Parameters:
        - pdf_path (str): The path to the input PDF file.
        - final_path (str): The directory where the highlighted PDFs will be saved.
        - text_to_highlight (list of str): List of texts to be highlighted in the PDF.
        - single_change_highlights (bool): If True, each highlighted text will be saved as a separate PDF;
                                        if False, all highlights will be saved in a single PDF.

        Returns:
        - list or None: If single_change_highlights is True, returns a list of saved PDF file paths with highlights;
                    otherwise, returns None.
        """
        if single_change_highlights:
            # If each highlighted text should be saved as a separate PDF
            link_list = []
            # Iterate over each text to highlight
            for i, text in enumerate(text_to_highlight):
                # Open the PDF file
                doc = fitz.open(pdf_path)
                # Iterate over each page in the PDF
                for j, page in enumerate(doc):
                    # Search for instances of the specified text on the current page
                    text_instances = page.search_for(text)
                    # Iterate over each instance of the text on the current page
                    for inst in text_instances:
                        # Add a highlight annotation to the instance
                        highlight = page.add_highlight_annot(inst)
                        highlight.update()
                    # Check if any instances were found on the page
                    if text_instances:
                        print("File saving..")
                        # Generate a unique file name based on the input PDF file name and index
                        file_name = pdf_path.split("/")[-1].replace(".", "").replace("pdf", "").replace("-", "").replace(
                            " ", "")
                        file_path = f"{final_path}//" + f"{file_name}_{i}.pdf"
                        print(file_path)
                        # Save the modified PDF with the highlight
                        doc.save(file_path)

                        # Create a link to the saved PDF with page number
                        file_n = file_path + f"#page={j + 1}"
                        link_list.append(file_n)
                        break # Break the page loop since we're saving each text separately
                # Close the opened PDF
                doc.close()

            # Return the list of saved PDF file paths with highlights
            return link_list
        else:
            # If all highlights should be saved in a single PDF
            doc = fitz.open(pdf_path)
            # Iterate over each text to highlight
            for i, text in enumerate(text_to_highlight):
                # print(text)
                # Iterate over each page in the PDF
                for j, page in enumerate(doc):
                    # Search for instances of the specified text on the current page
                    text_instances = page.search_for(text)
                    # Iterate over each instance of the text on the current page
                    for inst in text_instances:
                        # Add a highlight annotation to the instance
                        highlight = page.add_highlight_annot(inst)
                        highlight.update()
            print("File saving..")
            # Save the modified PDF with all highlights
            doc.save(f"{final_path}//" + "highlighted_all.pdf")
            # Close the opened PDF
            doc.close()
            # Generate the file path for the single PDF with all highlights
            file_n = f"{final_path}//" + f"highlighted_all.pdf"
            # Open the highlighted PDF in a web browser
            webbrowser.open(f'http://localhost:{8009}/{file_n}')

    def link_open(self,PORT:int=8009):
        """
        Start a simple HTTP server in the background to serve files.

        Parameters:
        - Handler (class): The request handler class (default is SimpleHTTPRequestHandler).
        - PORT (int): The port number on which the HTTP server will run (default is 8009).
        """
        Handler=http.server.SimpleHTTPRequestHandler
        def start_server(PORT=PORT):
            """
            Start the HTTP server.

            Parameters:
            - PORT (int): The port number on which the HTTP server will run.
            """
            with socketserver.TCPServer(("", PORT), Handler) as httpd:
                print("serving at port", PORT)
                httpd.serve_forever()

        # Start a background thread for the server
        server_thread = threading.Thread(target=start_server)
        server_thread.start()
        return server_thread

    def replace_newlines(self,text:str)-> str:
        """
        Replace consecutive sequences of three or more newlines with double newlines.

        Parameters:
        - text (str): The input text.

        Returns:
        - str: Text with consecutive sequences of three or more newlines replaced with double newlines.
        """
        return re.sub('(\n\s*){3,}', '\n\n', text)

    def download_docx(self,final_text:str, output_file_path:str="MSA_updated.docx") -> Document:
        """
        Create and download a DOCX file with specified content.

        Parameters:
        - final_text (str): The text content to be added to the DOCX file.
        - output_file_path (str): The path to save the generated DOCX file (default is "MSA_updated.docx").

        Returns:
        - Document: The Document object representing the created DOCX file.
        """
        # Open the template DOCX file
        f = open('New word tempate.docx', 'rb')
        document = Document(f)
        f.close()
        # doc=Document()
        # Set font style for the document
        style = document.styles['Normal']
        style.font.name = "Calibri"
        # Add the final text to the document, replacing "footer" if present
        document.add_paragraph(final_text.replace("footer", ""))
        # Save the document to the specified output file path
        document.save(output_file_path)
        # Return the Document object representing the created DOCX file
        return document

    def pdf_docx_converter(self,pdf_file_path:str, docx_file_dir:str)->str:
        """
        Convert a PDF file to a DOCX file.

        Parameters:
        - pdf_file_path (str): The path to the input PDF file.
        - docx_file_dir (str): The directory where the converted DOCX file will be saved.

        Returns:
        - str: The path to the converted DOCX file.
        """
        # Extract the name of the PDF file (excluding the extension)
        pdf_file_name = pdf_file_path.split('/')[-1].replace('.pdf', '')
        # Create the path for the output DOCX file
        docx_file_path = f"{docx_file_dir}//{pdf_file_name}.docx"
        # convert pdf to docx
        cv = Converter(pdf_file_path)
        cv.convert(docx_file_path)  # all pages by default
        cv.close()
        # Return the path to the converted DOCX file
        return docx_file_path

    def saved_updated_docx(self,updated_final_results:list[tuple], docx_file_path:str, new_docx_file_path:str)->list[str]:
        """
        Replace specified text in a DOCX file and save the updated document.

        Parameters:
        - updated_final_results (list of tuple): A list of tuples where each tuple contains the old and new text.
        - docx_file_path (str): The path to the input DOCX file.
        - new_docx_file_path (str): The path to save the updated DOCX file.

        Returns:
        - list: A list indicating whether each replacement was successful ("yes" or "No").
        """
        # Initialize a list to track update flags
        update_flag = []
        # Open the docx file
        doc = Document(docx_file_path)

        # Iterate through each tuple of old and new text
        for old, new in updated_final_results:
            # Initialize a flag to check if the replacement was successful
            check = False
            # Loop through the paragraphs and runs of the document
            for para in doc.paragraphs:
                if para.text.find(old) != -1:
                    # Replace the old text with the new text
                    para.text = para.text.replace(old, new)  # re.sub(old, new, run.text)
                    print("yes")
                    check = True
                    # Update the update_flag list
                    update_flag.append("yes")

            # If the replacement was not successful, update the flag list accordingly
            if not check:
                print("No")
                update_flag.append("No")
        # Save the updated document
        doc.save(new_docx_file_path)
        # Return the list of update flags
        return update_flag
    
    def saved_updated_docx_v2(self, updated_final_results, docx_file_path, new_docx_file_path):
        """
        Replaces specified text in a DOCX file and saves the updated document.

        Parameters:
        - updated_final_results (list): A list of tuples containing old and new text pairs.
        - docx_file_path (str): The path to the original DOCX file.
        - new_docx_file_path (str): The path to save the updated DOCX file.

        Returns:
        - list: A list of update flags indicating whether each specified text was found and replaced ('yes' or 'No').
        """
        # update_flag = []
        update_flag_dict={}
        # Open the docx file
        doc = Document(docx_file_path)

        for old, new in updated_final_results:
            check = False
            # Loop through the paragraphs and runs of the document
            for para in doc.paragraphs:
                # Check if the old text is present in the paragraph
                if re.sub('\s+', ' ', para.text.replace('\n', '')).strip().find(
                        re.sub('\s+', ' ', old.replace('\n', '')).strip()) != -1:
                    # Use regex to replace the text
                    para.text = re.sub('\s+', ' ', para.text.replace('\n', '')).strip().replace(
                        re.sub('\s+', ' ', old.replace('\n', '')).strip(), new)
                    print("yes")
                    check = True
                    # update_flag.append("yes")
                    update_flag_dict.update({old:"yes"})
            
            # Check if the document has tables
            if len(doc.tables) > 0:
                tables = doc.tables
                for i, table in enumerate(tables):
                    for row in table.rows:
                        for cell in row.cells:
                            # Check if the old text is present in the cell
                            if re.sub('\s+', ' ', cell.text.replace('\n', '')).strip().find(
                                    re.sub('\s+', ' ', old.replace('\n', '')).strip()) != -1:
                                cell.text = re.sub('\s+', ' ', cell.text.replace('\n', '')).strip().replace(
                                    re.sub('\s+', ' ', old.replace('\n', '')).strip(), new)
                                check = True
                                # update_flag.append("yes")
                                update_flag_dict.update({old:"yes"})
                                print(f"Table {i + 1}: yes")

            if not check:
                print("No")
                # update_flag.append("No")
                update_flag_dict.update({old:"No"})

        # Save the updated document
        doc.save(new_docx_file_path)

        return list(update_flag_dict.values())

    def updated_text_with_color(self,new_updated_file:str, updated_final_results:list[tuple]):
        """
        Highlight specific text in a DOCX file with a specified color and save the updated document.

        Parameters:
        - new_updated_file (str): The path to the input DOCX file to be updated.
        - updated_final_results (list of tuple): A list of tuples where each tuple contains the old and specific text.

        Returns:
        - None
        """
        # Load the document
        doc = Document(new_updated_file)
        new_color = RGBColor(0x42, 0x24, 0xE9)  # RGB color for blue

        # Iterate over the paragraphs in the document
        for paragraph in doc.paragraphs:
            # If the specific text is in the paragraph
            for old, specific_text in updated_final_results:
                if specific_text in paragraph.text:
                    # Split the paragraph text by the specific text
                    parts = paragraph.text.split(specific_text)

                    # Clear the paragraph text
                    paragraph.clear()

                    # Add the parts and the specific text with the new color to the paragraph
                    for i in range(len(parts)):
                        paragraph.add_run(parts[i])
                        if i < len(parts) - 1:
                            run = paragraph.add_run(specific_text)
                            run.font.color.rgb = new_color
                # Check if the document has tables
        if len(doc.tables) > 0:
            tables = doc.tables
            for old, specific_text in updated_final_results:
                for i, table in enumerate(tables):
                    for row in table.rows:
                        for cell in row.cells:
                            # Check if the old text is present in the cell
                            if re.sub('\s+', ' ', cell.text.replace('\n', '')).strip().find(
                                    re.sub('\s+', ' ', specific_text.replace('\n', '')).strip()) != -1:
                                # cell.font.color.rgb =new_color
                                cell.paragraphs[0].runs[0].font.color.rgb = new_color
        # Save the document
        f_name = new_updated_file.split("/")[-1]
        v2_file_path = f"version_2/{f_name}"
        doc.save(v2_file_path)

    def openai_api_calculate_cost(self,usage:dict, model:str="gpt-4-8k"):
        """
        Calculate the cost of OpenAI API usage based on the provided usage data and model.

        Parameters:
        - usage (dict): A dictionary containing the usage data, including 'prompt_tokens' and 'completion_tokens'.
        - model (str): The model for which the cost is to be calculated (default is "gpt-4-8k").

        Returns:
        - dict: A dictionary containing the calculated cost details.
        """
        # Define pricing information for different models
        pricing = {
            'gpt-3.5-turbo-4k': {
                'prompt': 0.0015,
                'completion': 0.002,
            },
            'gpt-3.5-turbo-16k': {
                'prompt': 0.003,
                'completion': 0.004,
            },
            'gpt-4-8k': {
                'prompt': 0.03,
                'completion': 0.06,
            },
            'gpt-4-32k': {
                'prompt': 0.06,
                'completion': 0.12,
            },
            'text-embedding-ada-002-v2': {
                'prompt': 0.0001,
                'completion': 0.0001,
            }
        }

        try:
            # Retrieve pricing information for the specified model
            model_pricing = pricing[model]
        except KeyError:
            raise ValueError("Invalid model specified")
        
        # Calculate the cost for prompt and completion tokens
        prompt_cost = usage['prompt_tokens'] * model_pricing['prompt'] / 1000
        completion_cost = usage['completion_tokens'] * model_pricing['completion'] / 1000
        # Calculate the total tokens and total cost
        total_tokens = usage['prompt_tokens'] + usage['completion_tokens']
        total_cost = prompt_cost + completion_cost

        # Print and return the calculated cost details
        print(
            f"\nTokens used:  {usage['prompt_tokens']:,} prompt + {usage['completion_tokens']:,} completion = {usage['total_tokens']:,} tokens")
        print(f"Total cost for {model}: ${total_cost:.4f}\n")

        return {
            "prompt_tokens": usage['prompt_tokens'],
            "prompt_cost": prompt_cost,
            "completion_tokens": usage['completion_tokens'],
            "completion_cost": completion_cost,
            "Total Tokens": total_tokens,
            "Total Cost": total_cost
        }

    def num_tokens_from_string(self,string: str, encoding_name: str = "cl100k_base") -> int:
        """
        Returns the number of tokens in a text string.

        Parameters:
        - string (str): The input text string.
        - encoding_name (str): The name of the encoding to use (default is "cl100k_base").

        Returns:
        - int: The number of tokens in the text string.
        """
        # Get the encoding based on the specified encoding name
        encoding = tiktoken.get_encoding(encoding_name)
        # Calculate the number of tokens in the text string using the encoding
        num_tokens = len(encoding.encode(string))
        # Return the number of tokens
        return num_tokens


    def docx_pdf(self,docx_file_path:str, pdf_file_name:str) -> str:
        """
        Convert a DOCX file to a PDF file.

        Parameters:
        - docx_file_path (str): The path to the input DOCX file.
        - pdf_file_name (str): The name of the output PDF file.

        Returns:
        - str: The name of the converted PDF file.
        """
        # Use the docx2pdf library to convert the DOCX file to PDF
        convert(docx_file_path, pdf_file_name)
        # Return the name of the converted PDF file
        return pdf_file_name
    def create_embedding_assistant(self, document_chunks):
        """
        Creates an instance of an embedding assistant using OpenAI's text embeddings.

        Parameters:
        - document_chunks (list): A list of text chunks to be used for creating embeddings.

        Returns:
        - vectorstore: An instance of Chroma containing text embeddings.

        Example usage:
        ```
        your_instance = YourClass()
        embeddings_assistant = your_instance.create_embedding_assistant(["text_chunk1", "text_chunk2"])
        ```

        Note:
        The method relies on the OpenAIEmbeddings and Chroma classes, which should be properly defined or imported.
        """
        # Initialize OpenAIEmbeddings with specified deployment, model, chunk_size, and max_retries
        embeddings = OpenAIEmbeddings(deployment="ada-text-embeddings-002",
                                      model="text-embedding-ada-002", chunk_size=1, max_retries=5)

        # Create a Chroma vectorstore from the provided document chunks using the embeddings
        vectorstore = Chroma.from_texts(texts=document_chunks, embedding=embeddings)

        # Return the created vectorstore
        return vectorstore
